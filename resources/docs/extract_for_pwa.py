"""
GCC Playbook PWA Content Extraction — v2
=========================================
Uses GCC_Playbook_v_0.9.docx as the MASTER source for Part I.
Adds Part II content from GCC_Playbook_Part_II_CLEAN.docx.
Adds Part III content from GCC_Playbook_Part_III.docx.
"""
import json
import os
import re
from docx import Document

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(os.path.dirname(DOCS_DIR), "..", "pwa", "public", "data")

MASTER_FILE = os.path.join(DOCS_DIR, "GCC_Playbook_v_0.9.docx")
PART2_FILE = os.path.join(DOCS_DIR, "GCC_Playbook_Part_II_CLEAN.docx")
PART3_FILE = os.path.join(DOCS_DIR, "GCC_Playbook_Part_III.docx")


def extract_chapters_from_doc(doc, part_label, chapter_prefix=""):
    """Extract structured chapter data from a docx, properly handling heading hierarchy."""
    chapters = []
    current_chapter = None
    current_section = None
    current_subsection = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        style_name = para.style.name if para.style else "Normal"
        
        if style_name == 'Heading 1':
            # Save previous chapter
            if current_chapter:
                chapters.append(current_chapter)
            
            current_chapter = {
                "id": f"{chapter_prefix}ch-{len(chapters)+1}",
                "title": text,
                "sections": [],
                "content": [],
            }
            current_section = None
            current_subsection = None
            
        elif style_name == 'Heading 2':
            if current_chapter is None:
                # H2 before any H1 — treat as intro section
                current_chapter = {
                    "id": f"{chapter_prefix}ch-intro",
                    "title": text,
                    "sections": [],
                    "content": [],
                }
                current_section = None
                current_subsection = None
                continue
            
            is_key = any(kw in text for kw in [
                'Chapter Summary', 'Things Every', 'Principles',
                'Ten Things', 'Ten Truths', 'Ten Community',
                'Ten KPI', 'Key Lessons'
            ])
            
            current_section = {
                "title": text,
                "content": [],
                "subsections": [],
                "isKeyTakeaway": is_key
            }
            current_chapter["sections"].append(current_section)
            current_subsection = None
            
        elif style_name == 'Heading 3':
            if current_section:
                current_subsection = {
                    "title": text,
                    "content": []
                }
                current_section["subsections"].append(current_subsection)
            elif current_chapter:
                # H3 without parent H2 — make it a section
                current_section = {
                    "title": text,
                    "content": [],
                    "subsections": [],
                    "isKeyTakeaway": False
                }
                current_chapter["sections"].append(current_section)
                current_subsection = None
        else:
            # Regular paragraph — add to deepest container
            if len(text) < 3:
                continue
            
            target = None
            if current_subsection is not None:
                target = current_subsection
            elif current_section is not None:
                target = current_section
            elif current_chapter is not None:
                target = current_chapter
            
            if target is not None and len(target.get("content", [])) < 25:
                target.setdefault("content", []).append(text)
    
    # Don't forget the last chapter
    if current_chapter:
        chapters.append(current_chapter)
    
    return chapters


def extract_glossary_from_doc(doc):
    """Extract glossary terms from a document that has a Key Terms section."""
    glossary = []
    in_glossary = False
    
    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name if para.style else "Normal"
        
        if text == "Key Terms and Definitions":
            in_glossary = True
            continue
        
        if in_glossary:
            # Stop at Subject Index or next major heading
            if text == "Subject Index" or (style_name == 'Heading 1' and 'Subject' in text):
                in_glossary = False
                continue
            
            if style_name.startswith('Heading'):
                continue  # Skip alphabet headers
            
            if not text:
                continue
            
            # Parse term: definition patterns
            for sep in ['—', ':', ' – ']:
                if sep in text:
                    parts = text.split(sep, 1)
                    if len(parts) == 2:
                        term = parts[0].strip()
                        definition = parts[1].strip()
                        if 3 < len(term) < 80 and len(definition) > 10:
                            glossary.append({"term": term, "definition": definition})
                            break
    
    return glossary


def extract_references_from_doc(doc, start_heading=""):
    """Extract references from a document."""
    refs = []
    in_refs = False
    ref_category = ""
    
    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name if para.style else "Normal"
        
        if any(kw in text for kw in ['VALIDATED REFERENCES', 'Primary Sources', 'How to Use This Playbook']):
            in_refs = True
            continue
        
        if in_refs:
            if style_name.startswith('Heading'):
                ref_category = text
                continue
            if text and not style_name.startswith('Heading'):
                refs.append({"category": ref_category, "text": text})
    
    return refs


def main():
    print("=" * 60)
    print("GCC Playbook PWA — Content Extraction v2")
    print("=" * 60)
    
    # ---- PART I: from v_0.9.docx ----
    print("\n[1/3] Extracting Part I from GCC_Playbook_v_0.9.docx...")
    master_doc = Document(MASTER_FILE)
    part1_chapters = extract_chapters_from_doc(master_doc, "Part I", "p1-")
    part1_glossary = extract_glossary_from_doc(master_doc)
    part1_refs = extract_references_from_doc(master_doc)
    print(f"  → {len(part1_chapters)} chapters, {len(part1_glossary)} glossary terms")
    
    # Separate appendices from main chapters
    # Chapters 1.x through 6.x are Part I main content
    # Appendices start with A.
    main_chapters = []
    appendix_chapters = []
    for ch in part1_chapters:
        title = ch["title"]
        if title.startswith("A.") or title.startswith("Appendix") or "Quick Navigation" in title:
            appendix_chapters.append(ch)
        elif "How to Use" in title:
            appendix_chapters.append(ch)
        else:
            main_chapters.append(ch)
    
    print(f"  → {len(main_chapters)} main chapters, {len(appendix_chapters)} appendix items")
    
    # ---- PART II: from Part_II_CLEAN.docx ----
    print("\n[2/3] Extracting Part II from GCC_Playbook_Part_II_CLEAN.docx...")
    part2_doc = Document(PART2_FILE)
    part2_chapters = extract_chapters_from_doc(part2_doc, "Part II", "p2-")
    part2_glossary = extract_glossary_from_doc(part2_doc)
    part2_refs = extract_references_from_doc(part2_doc)
    print(f"  → {len(part2_chapters)} chapters, {len(part2_glossary)} glossary terms")
    
    # ---- PART III: from Part_III.docx ----
    print("\n[3/3] Extracting Part III from GCC_Playbook_Part_III.docx...")
    part3_doc = Document(PART3_FILE)
    part3_chapters = extract_chapters_from_doc(part3_doc, "Part III", "p3-")
    part3_refs = extract_references_from_doc(part3_doc)
    print(f"  → {len(part3_chapters)} chapters")
    
    # ---- Combine glossary ----
    all_glossary = part1_glossary + part2_glossary
    # Deduplicate
    seen_terms = set()
    unique_glossary = []
    for g in all_glossary:
        key = g["term"].lower()
        if key not in seen_terms:
            seen_terms.add(key)
            unique_glossary.append(g)
    unique_glossary.sort(key=lambda x: x["term"].lower())
    
    # ---- Combine references ----
    all_refs = []
    for r in part1_refs + part2_refs + part3_refs:
        if r["text"] not in [x["text"] for x in all_refs]:
            all_refs.append(r)
    
    # ---- Build output ----
    data = {
        "title": "The Complete India GCC Reference 2026–2030",
        "subtitle": "Landscape | Maturity Models | Operating Models | Challenges | KPIs | Community | AI | Deep Tech | M&A | Finance",
        "author": "Curated by Kalilur Rahman",
        "parts": {
            "part1": {
                "title": "Part I: India's GCC Landscape",
                "subtitle": "Landscape, Maturity, Operating Models, Challenges, KPIs, Community & Ecosystem",
                "chapters": main_chapters
            },
            "part2": {
                "title": "Part II: GCC by Scale, Stage & Leadership",
                "subtitle": "Sizing, Setup, Lifecycle, Operating Models, Case Studies, Who's Who",
                "chapters": part2_chapters
            },
            "part3": {
                "title": "Part III: The Frontier",
                "subtitle": "AI, Deep Tech, M&A, Finance, Legal, Innovation, ESG, Career, 2030 Scenarios",
                "chapters": part3_chapters
            },
            "appendices": {
                "title": "Appendices",
                "subtitle": "Decision Matrices, Scorecards, Checklists, Templates, Resource Directory",
                "chapters": appendix_chapters
            }
        },
        "glossary": unique_glossary,
        "references": [r["text"] for r in all_refs[:100]],
        "stats": {
            "totalChapters": len(main_chapters) + len(part2_chapters) + len(part3_chapters),
            "part1Chapters": len(main_chapters),
            "part2Chapters": len(part2_chapters),
            "part3Chapters": len(part3_chapters),
            "appendices": len(appendix_chapters),
            "glossaryTerms": len(unique_glossary),
            "references": len(all_refs)
        }
    }
    
    # Write output
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    output_path = os.path.join(OUTPUT_DIR, "content.json")
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    
    print(f"\n{'=' * 60}")
    print(f"EXTRACTION COMPLETE")
    print(f"{'=' * 60}")
    print(f"Output: {output_path}")
    print(f"Stats:")
    for k, v in data["stats"].items():
        print(f"  {k}: {v}")
    
    # Print chapter list for verification
    print(f"\n--- Part I Chapters ---")
    for ch in main_chapters:
        sections = len(ch.get("sections", []))
        content = len(ch.get("content", []))
        print(f"  {ch['title'][:70]:<72} ({sections} sections, {content} intro paras)")
    
    print(f"\n--- Part II Chapters ---")
    for ch in part2_chapters:
        sections = len(ch.get("sections", []))
        content = len(ch.get("content", []))
        print(f"  {ch['title'][:70]:<72} ({sections} sections, {content} intro paras)")
    
    print(f"\n--- Part III Chapters ---")
    for ch in part3_chapters:
        sections = len(ch.get("sections", []))
        content = len(ch.get("content", []))
        print(f"  {ch['title'][:70]:<72} ({sections} sections, {content} intro paras)")


if __name__ == "__main__":
    main()
